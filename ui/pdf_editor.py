import logging
import io
import os
from PySide6.QtWidgets import (QGraphicsView, QGraphicsScene, QVBoxLayout,
                               QPushButton, QWidget, QHBoxLayout, QFileDialog,
                               QGraphicsTextItem, QGraphicsItem, QGraphicsObject,
                               QGraphicsPixmapItem,
                               QDialog, QComboBox, QDoubleSpinBox, QApplication,
                               QLabel, QMenu, QColorDialog)
from PySide6.QtCore import Qt, Signal, QRectF, QPointF, QSettings
from PySide6.QtGui import (QPixmap, QPen, QBrush, QPainter, QColor, QFont,
                              QTextCursor, QTextBlockFormat, QTextCharFormat,
                              QShortcut, QKeySequence)
from PySide6.QtPrintSupport import QPrinter
from PIL import Image as PILImage
from core.font_utils import WordFontSizeAdapter

logger = logging.getLogger(__name__)

A4_W = 210
A4_H = 297
FONT_SCALE = 0.234375  # 0.375 * 10/16 — combo "16" matches what "10" used to show


class PdfGraphicsView(QGraphicsView):
    file_dropped = Signal(str)
    page_clicked = Signal(float, float)
    image_dropped = Signal(str)
    delete_pressed = Signal()
    paste_triggered = Signal()
    copy_triggered = Signal()

    def __init__(self, scene, parent=None):
        super().__init__(scene, parent)
        self.setAcceptDrops(True)
        self.setRenderHint(QPainter.RenderHint.Antialiasing)
        self.setRenderHint(QPainter.RenderHint.SmoothPixmapTransform)
        self.setDragMode(QGraphicsView.NoDrag)
        self._is_panning = False
        self._pan_start = None

    def wheelEvent(self, event):
        self.verticalScrollBar().setValue(
            self.verticalScrollBar().value() - event.angleDelta().y())
        event.accept()

    def mousePressEvent(self, event):
        if event.button() == Qt.MiddleButton:
            self._is_panning = True
            self._pan_start = event.pos()
            self.setCursor(Qt.ClosedHandCursor)
            event.accept()
            return
        if event.button() == Qt.LeftButton:
            p = self.mapToScene(event.pos())
            self.page_clicked.emit(p.x(), p.y())
        super().mousePressEvent(event)

    def mouseMoveEvent(self, event):
        if self._is_panning and self._pan_start is not None:
            delta = event.pos() - self._pan_start
            self._pan_start = event.pos()
            self.horizontalScrollBar().setValue(
                self.horizontalScrollBar().value() - delta.x())
            self.verticalScrollBar().setValue(
                self.verticalScrollBar().value() - delta.y())
            event.accept()
            return
        super().mouseMoveEvent(event)

    def mouseReleaseEvent(self, event):
        if event.button() == Qt.MiddleButton and self._is_panning:
            self._is_panning = False
            self._pan_start = None
            self.setCursor(Qt.ArrowCursor)
            event.accept()
            return
        super().mouseReleaseEvent(event)

    def dragEnterEvent(self, event):
        if event.mimeData().hasUrls():
            event.acceptProposedAction()

    def dragMoveEvent(self, event):
        if event.mimeData().hasUrls():
            event.acceptProposedAction()

    def dropEvent(self, event):
        for url in event.mimeData().urls():
            path = url.toLocalFile()
            if path.lower().endswith('.pdf'):
                self.file_dropped.emit(path)
            else:
                ext = os.path.splitext(path)[1].lower()
                if ext in ('.png', '.jpg', '.jpeg', '.bmp', '.gif', '.tiff', '.webp'):
                    self.image_dropped.emit(path)
        event.acceptProposedAction()

    def keyPressEvent(self, event):
        if event.key() in (Qt.Key_Delete, Qt.Key_Backspace):
            self.delete_pressed.emit()
        elif event.key() == Qt.Key_V and event.modifiers() == Qt.ControlModifier:
            self.paste_triggered.emit()
        elif event.key() == Qt.Key_C and event.modifiers() == Qt.ControlModifier:
            self.copy_triggered.emit()
        super().keyPressEvent(event)


class ConstrainedTextItem(QGraphicsTextItem):
    def __init__(self, editor):
        super().__init__()
        self._editor = editor

    def itemChange(self, change, value):
        if change == QGraphicsItem.ItemPositionChange and self.scene():
            m = self._editor._margin_mm
            left = 2 + m
            right = A4_W + 2 - m
            w = self.boundingRect().width() * self.scale()
            nx = value.x()
            if nx < left:
                nx = left
            if nx + w > right:
                nx = right - w
            value.setX(nx)
        return super().itemChange(change, value)


class ResizableImageItem(QGraphicsObject):
    HANDLE_R = 2
    LINE_LEN = 4
    MIN_SIZE = 20

    def __init__(self, pixmap):
        super().__init__()
        self._pix = pixmap
        self._source_rect = QRectF(0, 0, pixmap.width(), pixmap.height())
        self._rect = QRectF(0, 0, pixmap.width(), pixmap.height())
        self._drag_handle = -1
        self._drag_rect = None
        self._drag_start = None
        self._crop_mode = False
        self.setFlag(QGraphicsItem.ItemIsMovable, True)
        self.setFlag(QGraphicsItem.ItemIsSelectable, True)

    def boundingRect(self):
        m = self.HANDLE_R + 1
        return self._rect.adjusted(-m, -m, m, m)

    def paint(self, painter, option, widget):
        painter.setRenderHint(QPainter.SmoothPixmapTransform)
        painter.drawPixmap(self._rect, self._pix, self._source_rect)
        painter.setBrush(Qt.NoBrush)
        painter.setPen(QPen(QColor(0, 120, 215), 1.5))
        painter.drawRect(self._rect)
        if self.isSelected():
            painter.setBrush(Qt.white)
            painter.setPen(QPen(QColor(0, 120, 215), 1.5))
            for pt in self._corner_handles():
                painter.drawEllipse(pt, self.HANDLE_R, self.HANDLE_R)
        if self._crop_mode:
            painter.setBrush(Qt.white)
            painter.setPen(QPen(QColor(0, 120, 215), 1.5))
            r = self._rect
            cx = (r.left() + r.right()) / 2
            cy = (r.top() + r.bottom()) / 2
            painter.drawRect(QRectF(cx - self.LINE_LEN, r.top() - 1.5, self.LINE_LEN * 2, 3))
            painter.drawRect(QRectF(r.right() - 1.5, cy - self.LINE_LEN, 3, self.LINE_LEN * 2))
            painter.drawRect(QRectF(cx - self.LINE_LEN, r.bottom() - 1.5, self.LINE_LEN * 2, 3))
            painter.drawRect(QRectF(r.left() - 1.5, cy - self.LINE_LEN, 3, self.LINE_LEN * 2))

    def _corner_handles(self):
        return [self._rect.topLeft(), self._rect.topRight(),
                self._rect.bottomLeft(), self._rect.bottomRight()]

    def _edge_positions(self):
        r = self._rect
        cx = (r.left() + r.right()) / 2
        cy = (r.top() + r.bottom()) / 2
        return [QPointF(cx, r.top()), QPointF(r.right(), cy),
                QPointF(cx, r.bottom()), QPointF(r.left(), cy)]

    def _handle_at(self, pos):
        for i, pt in enumerate(self._corner_handles()):
            if (pos - pt).manhattanLength() <= self.HANDLE_R * 4:
                return i
        for i, pt in enumerate(self._edge_positions(), start=4):
            if (pos - pt).manhattanLength() <= 8:
                return i
        return -1

    def mousePressEvent(self, event):
        self._drag_handle = self._handle_at(event.pos())
        if self._drag_handle >= 0:
            self._drag_rect = QRectF(self._rect)
            self._drag_start = event.pos()
            event.accept()
            return
        super().mousePressEvent(event)
        event.accept()

    def mouseMoveEvent(self, event):
        if self._drag_handle >= 0:
            d = event.pos() - self._drag_start
            r = QRectF(self._drag_rect)
            h = self._drag_handle
            if h < 4:  # corner → proportional
                ar = self._drag_rect.width() / max(self._drag_rect.height(), 1)
                ox = self._drag_rect.left() if h in (1, 3) else self._drag_rect.right()
                oy = self._drag_rect.top() if h in (2, 3) else self._drag_rect.bottom()
                sx = 1 if h in (1, 3) else -1
                sy = 1 if h in (2, 3) else -1
                dx = self._drag_rect.width() * sx + d.x()
                dy = self._drag_rect.height() * sy + d.y()
                adx, ady = abs(dx), abs(dy)
                if adx / max(ady, 1) > ar:
                    ady = adx / ar
                else:
                    adx = ady * ar
                w, hh = adx, ady
                nx = ox - w if sx < 0 else ox
                ny = oy - hh if sy < 0 else oy
                r = QRectF(nx, ny, w, hh)
            else:  # edge → single axis
                if h == 4:    # top
                    r.setTop(r.top() + d.y())
                elif h == 5:  # right
                    r.setRight(r.right() + d.x())
                elif h == 6:  # bottom
                    r.setBottom(r.bottom() + d.y())
                elif h == 7:  # left
                    r.setLeft(r.left() + d.x())
            if r.width() >= self.MIN_SIZE and r.height() >= self.MIN_SIZE:
                if self._crop_mode and h >= 4:
                    self._apply_crop(r)
                else:
                    self._rect = r
                    self.prepareGeometryChange()
                    self.update()
            event.accept()
            return
        super().mouseMoveEvent(event)

    def mouseReleaseEvent(self, event):
        self._drag_handle = -1
        self._drag_rect = None
        super().mouseReleaseEvent(event)

    def contextMenuEvent(self, event):
        menu = QMenu()
        crop_action = menu.addAction("✂️ قص" if not self._crop_mode else "✂️ إلغاء القص")
        copy_action = menu.addAction("📋 نسخ")
        menu.addSeparator()
        raise_action = menu.addAction("⬆️ رفع للأمام")
        lower_action = menu.addAction("⬇️ إرسال للخلف")
        chosen = menu.exec(event.screenPos())
        if chosen == crop_action:
            self._crop_mode = not self._crop_mode
            self.update()
            logger.info("وضع القص: %s", "مفعل" if self._crop_mode else "معطل")
        elif chosen == copy_action:
            QApplication.clipboard().setPixmap(self._pix)
            logger.info("تم نسخ الصورة")
        elif chosen == raise_action:
            self.setZValue(self.zValue() + 1)
            logger.info("رفع الصورة للأمام (z=%d)", self.zValue())
        elif chosen == lower_action:
            self.setZValue(max(0, self.zValue() - 1))
            logger.info("إرسال الصورة للخلف (z=%d)", self.zValue())

    def _apply_crop(self, new_rect):
        r = self._rect
        sr = self._source_rect
        scale_x = sr.width() / r.width() if r.width() > 0 else 1
        scale_y = sr.height() / r.height() if r.height() > 0 else 1
        new_sr = QRectF(
            sr.left() + (new_rect.left() - r.left()) * scale_x,
            sr.top() + (new_rect.top() - r.top()) * scale_y,
            new_rect.width() * scale_x,
            new_rect.height() * scale_y
        )
        pw, ph = self._pix.width(), self._pix.height()
        if new_sr.left() < 0:
            new_sr.setLeft(0)
        if new_sr.top() < 0:
            new_sr.setTop(0)
        if new_sr.right() > pw:
            new_sr.setRight(pw)
        if new_sr.bottom() > ph:
            new_sr.setBottom(ph)
        self._source_rect = new_sr
        self._rect = QRectF(new_rect)
        self.prepareGeometryChange()
        self.update()
        logger.info("تم قص الصورة")


class PdfEditor(QWidget):
    go_back = Signal()

    def __init__(self, parent=None):
        super().__init__(parent)
        self._num_pages = 1
        self._pdf_path = None
        self._text_regions = []
        self._current_editor = None
        self._margin_mm = 3.5
        self._undo_stack = []
        self._copy_pos = None
        self._copy_rect = None
        self._landscape = False
        self.setLayoutDirection(Qt.RightToLeft)
        self._build_ui()

    def _build_ui(self):
        layout = QVBoxLayout(self)
        layout.setContentsMargins(0, 0, 0, 0)

        btn_bar1 = QHBoxLayout()
        btn_bar2 = QHBoxLayout()

        # --- Row 1: file / navigation ---
        btn_back = QPushButton("← رجوع")
        btn_back.clicked.connect(self.go_back)
        btn_open = QPushButton("📂 فتح PDF")
        btn_open.clicked.connect(self._open_pdf)
        btn_img = QPushButton("🖼️ صورة")
        btn_img.clicked.connect(self._add_image_from_file)
        btn_print = QPushButton("🖨️ طباعة")
        btn_print.clicked.connect(self._print_pdf)
        btn_save = QPushButton("💾 حفظ")
        btn_save.setToolTip("حفظ المستند")
        save_menu = QMenu(self)
        save_menu.addAction("PDF", self._save_as_pdf)
        save_menu.addAction("Word", self._save_as_docx)
        btn_save.setMenu(save_menu)

        btn_del_file = QPushButton("🗑️ حذف الملف")
        btn_del_file.clicked.connect(self._delete_all_content)

        btn_orient = QPushButton("🔄 طولي")
        btn_orient.clicked.connect(self._toggle_orientation)

        self._page_combo = QComboBox()
        self._page_combo.setEditable(True)
        self._page_combo.setFixedWidth(50)
        self._page_combo.lineEdit().returnPressed.connect(self._page_combo_changed)
        self._page_total_label = QLabel("/ 1")
        btn_zin = QPushButton("🔍+")
        btn_zin.clicked.connect(lambda: (self._zoom_in(), self._save_zoom()))
        btn_zout = QPushButton("🔍−")
        btn_zout.clicked.connect(lambda: (self._zoom_out(), self._save_zoom()))

        btn_bar1.addWidget(btn_back)
        btn_bar1.addWidget(btn_open)
        btn_bar1.addWidget(btn_img)
        btn_bar1.addWidget(btn_print)
        btn_bar1.addWidget(btn_save)
        btn_bar1.addWidget(btn_del_file)
        btn_bar1.addWidget(btn_orient)
        btn_bar1.addStretch()
        btn_bar1.addWidget(QLabel("صفحة:"))
        btn_bar1.addWidget(self._page_combo)
        btn_bar1.addWidget(self._page_total_label)
        btn_bar1.addWidget(QLabel("هامش:"))
        self._margin_spin = QDoubleSpinBox()
        self._margin_spin.setRange(0, 30)
        self._margin_spin.setSingleStep(0.5)
        self._margin_spin.setValue(self._margin_mm)
        self._margin_spin.setFixedWidth(80)
        self._margin_spin.setDecimals(1)
        self._margin_spin.valueChanged.connect(self._margin_spin_changed)
        btn_bar1.addWidget(self._margin_spin)
        btn_bar1.addWidget(btn_zout)
        btn_bar1.addWidget(btn_zin)

        # --- Row 2: text formatting ---
        self.font_size_combo = QComboBox()
        self.font_size_combo.setEditable(True)
        for sz in WordFontSizeAdapter.WORD_SIZES:
            self.font_size_combo.addItem(str(sz), sz)
        self.font_size_combo.setCurrentText("12")
        self.font_size_combo.setFixedWidth(70)
        self.font_size_combo.lineEdit().setAlignment(Qt.AlignLeft)
        self.font_size_combo.currentIndexChanged.connect(self._font_size_combo_changed)
        self.font_size_combo.lineEdit().returnPressed.connect(self._font_size_combo_changed)
        btn_font_down = QPushButton("▼")
        btn_font_down.setFixedWidth(28)
        btn_font_down.clicked.connect(self._font_size_down)
        btn_font_up = QPushButton("▲")
        btn_font_up.setFixedWidth(28)
        btn_font_up.clicked.connect(self._font_size_up)

        self._bold_btn = QPushButton("B")
        self._bold_btn.setCheckable(True)
        self._bold_btn.setFixedWidth(32)
        self._bold_btn.setStyleSheet(
            "QPushButton { font-weight: bold; font-size: 14pt; }"
            "QPushButton:checked { background-color: #4a90d9; color: white; }"
        )
        self._bold_btn.clicked.connect(self._toggle_bold)

        self._color_btn = QPushButton("A")
        self._color_btn.setFixedWidth(32)
        self._color_btn.setToolTip("تغيير لون النص")
        self._color_btn.setStyleSheet(
            "QPushButton { font-weight: bold; font-size: 14pt; "
            "border-style: solid; border-width: 1px 1px 6px 1px; "
            "border-color: gray gray red gray; }")
        self._color_btn.clicked.connect(self._pick_color)

        self._hl_btn = QPushButton("🖍")
        self._hl_btn.setFixedWidth(32)
        self._hl_btn.setToolTip("تحديد بخلفية ملونة")
        self._hl_btn.setStyleSheet(
            "QPushButton { font-size: 14pt; }")
        self._hl_btn.clicked.connect(self._pick_highlight)

        self._align_left_btn = QPushButton("≡")
        self._align_left_btn.setFixedWidth(32)
        self._align_left_btn.setToolTip("محاذاة لليسار")
        self._align_left_btn.clicked.connect(lambda: self._set_alignment(Qt.AlignLeft))
        self._align_center_btn = QPushButton("☰")
        self._align_center_btn.setFixedWidth(32)
        self._align_center_btn.setToolTip("توسيط")
        self._align_center_btn.clicked.connect(lambda: self._set_alignment(Qt.AlignCenter))
        self._align_right_btn = QPushButton("≡")
        self._align_right_btn.setFixedWidth(32)
        self._align_right_btn.setToolTip("محاذاة لليمين")
        self._align_right_btn.clicked.connect(lambda: self._set_alignment(Qt.AlignRight))

        self._num_btn = QPushButton("١٢٣")
        self._num_btn.setCheckable(True)
        self._num_btn.setFixedWidth(36)
        self._num_btn.setToolTip("تحويل الأرقام عربي/إنجليزي")
        self._num_btn.setStyleSheet(
            "QPushButton:checked { background-color: #4a90d9; color: white; }"
        )
        self._num_btn.clicked.connect(self._toggle_numbers)

        btn_bar2.addWidget(self.font_size_combo)
        btn_bar2.addWidget(btn_font_down)
        btn_bar2.addWidget(btn_font_up)
        btn_bar2.addWidget(self._bold_btn)
        btn_bar2.addWidget(self._color_btn)
        btn_bar2.addWidget(self._hl_btn)
        btn_bar2.addWidget(self._align_left_btn)
        btn_bar2.addWidget(self._align_center_btn)
        btn_bar2.addWidget(self._align_right_btn)
        btn_bar2.addWidget(self._num_btn)
        btn_bar2.addStretch()

        layout.addLayout(btn_bar1)
        layout.addLayout(btn_bar2)

        self.scene = QGraphicsScene(self)
        self.scene.setSceneRect(0, 0, A4_W, A4_H)
        self.scene.setBackgroundBrush(QColor(200, 200, 200))
        self._draw_page_area(0)
        self.view = PdfGraphicsView(self.scene, self)
        self.view.file_dropped.connect(self._load_pdf)
        self.view.image_dropped.connect(self._add_image)
        self.view.page_clicked.connect(self._on_page_clicked)
        self.scene.selectionChanged.connect(self._on_sel_changed)
        self.view.verticalScrollBar().valueChanged.connect(self._update_page_label)
        self._undo_shortcut = QShortcut(QKeySequence.Undo, self)
        self._undo_shortcut.activated.connect(self._undo_last)
        self.view.delete_pressed.connect(self._delete_selected)
        self.view.paste_triggered.connect(self._paste_image)
        self.view.copy_triggered.connect(self._copy_selected)
        layout.addWidget(self.view)
        self._apply_default_zoom()

    def _get_default_zoom(self):
        settings = QSettings("ورشة طباعة", "PdfEditor")
        return settings.value("defaultZoom", 2.2, type=float)

    def _apply_default_zoom(self):
        zoom = self._get_default_zoom()
        self.view.scale(zoom, zoom)
        logger.info("تم تطبيق التكبير الافتراضي: %.2f", zoom)

    def _save_zoom(self):
        t = self.view.transform()
        zoom = t.m11()
        settings = QSettings("ورشة طباعة", "PdfEditor")
        settings.setValue("defaultZoom", zoom)

    def _draw_page_area(self, page_index):
        y0 = page_index * A4_H
        self.scene.addRect(2, y0 + 2, A4_W, A4_H, QPen(QColor(160, 160, 160)), QBrush(Qt.white)).setZValue(-1)

    def _open_pdf(self):
        path, _ = QFileDialog.getOpenFileName(self, "فتح ملف PDF", "", "PDF (*.pdf)")
        if path:
            self._load_pdf(path)

    def _load_pdf(self, path):
        import fitz
        doc = fitz.open(path)
        self._pdf_path = path
        self._text_regions.clear()
        self.scene.clear()
        self._num_pages = len(doc)
        self._page_combo.blockSignals(True)
        self._page_combo.clear()
        for p in range(1, self._num_pages + 1):
            self._page_combo.addItem(str(p), p)
        self._page_combo.setCurrentText("1")
        self._page_combo.blockSignals(False)
        self._page_total_label.setText(f"/ {self._num_pages}")
        self.scene.setSceneRect(0, 0, A4_W, self._num_pages * A4_H)

        for i in range(self._num_pages):
            self._draw_page_area(i)
            page = doc[i]
            pw = page.rect.width
            ph = page.rect.height
            y0 = i * A4_H

            zoom = 200 / 72.0
            pix = page.get_pixmap(matrix=fitz.Matrix(zoom, zoom))
            scale = min(A4_W / (pw * zoom), A4_H / (ph * zoom))

            # Extract text regions and store for click-to-edit
            for block in page.get_text("dict", flags=fitz.TEXT_PRESERVE_WHITESPACE)["blocks"]:
                if block.get("type") != 0:
                    continue
                for line in block.get("lines", []):
                    spans = line.get("spans", [])
                    if not spans:
                        continue
                    full_text = " ".join(s.get("text", "").strip() for s in spans if s.get("text", "").strip())
                    if not full_text:
                        continue
                    x0 = min(s["bbox"][0] for s in spans)
                    y0b = min(s["bbox"][1] for s in spans)
                    x1 = max(s["bbox"][2] for s in spans)
                    y1b = max(s["bbox"][3] for s in spans)
                    sx = x0 * zoom * scale
                    sy = y0 + y0b * zoom * scale
                    sw = (x1 - x0) * zoom * scale
                    sh = (y1b - y0b) * zoom * scale
                    first = spans[0]
                    self._text_regions.append((
                        QRectF(sx, sy, sw, sh),
                        full_text,
                        first.get("font", "Arial"),
                        first.get("color", 0),
                        first.get("size", 10)))

            # Render page background
            img = PILImage.frombytes("RGB", [pix.width, pix.height], pix.samples)
            buf = io.BytesIO()
            img.save(buf, format="PNG")
            qpix = QPixmap()
            qpix.loadFromData(buf.getvalue())
            bg_img = self.scene.addPixmap(qpix)
            bg_img.setPos(0, y0)
            bg_img.setScale(scale)
            bg_img.setZValue(0)

            # Extract embedded images from PDF page
            try:
                for img_ref in page.get_images(full=True):
                    xref = img_ref[0]
                    base = doc.extract_image(xref)
                    if not base or "image" not in base:
                        continue
                    img_bytes = base["image"]
                    img_qpix = QPixmap()
                    if not img_qpix.loadFromData(img_bytes):
                        continue
                    rects = page.get_image_rects(xref)
                    for r in rects:
                        sx = r.x0 * zoom * scale
                        sy = y0 + r.y0 * zoom * scale
                        sw = (r.x1 - r.x0) * zoom * scale
                        sh = (r.y1 - r.y0) * zoom * scale
                        if sw < 5 or sh < 5:
                            continue
                        img_item = ResizableImageItem(img_qpix)
                        img_item._rect = QRectF(0, 0, sw, sh)
                        img_item._source_rect = QRectF(0, 0, img_qpix.width(), img_qpix.height())
                        img_item.setPos(sx, sy)
                        img_item.setZValue(1)
                        self.scene.addItem(img_item)
                        # Cover original image in page background
                        cover = self.scene.addRect(sx, sy, sw, sh, QPen(Qt.NoPen), QBrush(Qt.white))
                        cover.setZValue(0)
                        logger.debug("استخراج صورة من PDF: %s (%.0f×%.0f)", base.get("ext", ""), sw, sh)
            except Exception:
                logger.debug("فشل استخراج الصور من الصفحة %d", i + 1)

            QApplication.processEvents()

        self.view.resetTransform()
        zoom = self._get_default_zoom()
        self.view.scale(zoom, zoom)
        self.view.centerOn(A4_W / 2, A4_H / 2)
        doc.close()
        logger.info("تم فتح PDF (%d صفحة) – zoom=%.2f", self._num_pages, zoom)

    def _toggle_orientation(self):
        global A4_W, A4_H
        old_w, old_h = A4_W, A4_H
        A4_W, A4_H = A4_H, A4_W
        self._landscape = not self._landscape
        sx_ratio = A4_W / old_w
        sy_ratio = A4_H / old_h
        # Save all scene content
        saved_texts = []
        saved_images = []
        saved_pages = []
        for item in self.scene.items():
            if isinstance(item, QGraphicsTextItem):
                p = item.pos()
                saved_texts.append({
                    'text': item.toPlainText(),
                    'font': QFont(item.font()),
                    'color': QColor(item.defaultTextColor()),
                    'x': p.x(), 'y': p.y(), 'rot': item.rotation(),
                })
            elif isinstance(item, ResizableImageItem):
                p = item.pos()
                saved_images.append({
                    'pixmap': item._pix,
                    'x': p.x(), 'y': p.y(),
                    'rect': QRectF(item._rect),
                    'source': QRectF(item._source_rect),
                    'z': item.zValue(),
                })
            elif isinstance(item, QGraphicsPixmapItem):
                p = item.pos()
                saved_pages.append({
                    'pixmap': item.pixmap(),
                    'x': p.x(), 'y': p.y(),
                    'scale': item.scale(),
                })
        saved_regions = [(QRectF(r), t, f, c, s) for r, t, f, c, s in self._text_regions]
        saved_path = self._pdf_path
        self._current_editor = None
        # Rebuild scene with new orientation
        self.scene.clear()
        for pg in range(self._num_pages):
            self._draw_page_area(pg)
        self.scene.setSceneRect(0, 0, A4_W, self._num_pages * A4_H)
        # Restore PDF page backgrounds with mapped positions
        for data in saved_pages:
            bg = self.scene.addPixmap(data['pixmap'])
            bg.setPos(data['x'] * sx_ratio, data['y'] * sy_ratio)
            bg.setScale(data['scale'])
            bg.setZValue(0)
        # Restore text regions with mapped coordinates
        new_regions = []
        for rect, text, font, color, size in saved_regions:
            new_regions.append((
                QRectF(rect.x() * sx_ratio, rect.y() * sy_ratio,
                       rect.width() * sx_ratio, rect.height() * sy_ratio),
                text, font, color, size))
        self._text_regions[:] = new_regions
        # Restore text items with full-width white background rects
        for data in saved_texts:
            item = ConstrainedTextItem(self)
            item.setPlainText(data['text'])
            item.setFont(data['font'])
            item.setDefaultTextColor(data['color'])
            item.setPos(data['x'] * sx_ratio, data['y'] * sy_ratio)
            item.setRotation(data['rot'])
            item.setZValue(2)
            self.scene.addItem(item)
            m = self._margin_mm
            pad = 0.3
            wr = self.scene.addRect(
                2 + m, item.pos().y() - pad,
                A4_W - 2 * m, item.boundingRect().height() + 2 * pad,
                QPen(Qt.NoPen), QBrush(Qt.white))
            wr.setZValue(1)
            wr.stackBefore(item)
        # Restore images
        for data in saved_images:
            img = ResizableImageItem(data['pixmap'])
            img.setPos(data['x'] * sx_ratio, data['y'] * sy_ratio)
            img._rect = QRectF(
                data['rect'].x() * sx_ratio, data['rect'].y() * sy_ratio,
                data['rect'].width() * sx_ratio, data['rect'].height() * sy_ratio)
            img._source_rect = QRectF(data['source'])
            img.setZValue(data.get('z', 5))
            self.scene.addItem(img)
        self.view.resetTransform()
        zoom = self._get_default_zoom()
        self.view.scale(zoom, zoom)
        self.view.centerOn(A4_W / 2, A4_H / 2)
        btn = self.sender()
        if btn:
            btn.setText("🔄 عرضي" if self._landscape else "🔄 طولي")
        logger.info("تم تغيير الاتجاه إلى %s (مع حفظ المحتوى)", "عرضي" if self._landscape else "طولي")

    def _delete_all_content(self):
        self._text_regions.clear()
        self._current_editor = None
        self.scene.clear()
        self._draw_page_area(0)
        self._num_pages = 1
        self._pdf_path = None
        self._undo_stack.clear()
        self._page_combo.blockSignals(True)
        self._page_combo.clear()
        self._page_combo.addItem("1", 1)
        self._page_combo.setCurrentText("1")
        self._page_combo.blockSignals(False)
        self._page_total_label.setText("/ 1")
        self.scene.setSceneRect(0, 0, A4_W, A4_H)
        logger.info("تم حذف جميع المحتوى")

    def _on_page_clicked(self, sx, sy):
        from PySide6.QtCore import QPointF
        items = self.scene.items(QPointF(sx, sy))
        if items and isinstance(items[0], ResizableImageItem):
            return
        for item in items:
            if isinstance(item, QGraphicsTextItem):
                item.setTextInteractionFlags(Qt.TextEditorInteraction)
                self.scene.setFocusItem(item)
                self._current_editor = item
                return
        for i, (rect, text, pdf_font, color_int, font_size) in enumerate(self._text_regions):
            if rect.contains(sx, sy):
                self._create_editor(rect, text, pdf_font, color_int, font_size)
                self._text_regions.pop(i)
                return
        self.scene.clearSelection()
        if self._current_editor is not None:
            self._current_editor.setTextInteractionFlags(Qt.NoTextInteraction)
            self._current_editor.clearFocus()
        self._current_editor = None

    def _create_editor(self, scene_rect, text, pdf_font, color_int, font_size):
        pad = 0.3
        m = self._margin_mm
        wr = self.scene.addRect(
            2 + m, scene_rect.y() - pad,
            A4_W - 2 * m, scene_rect.height() + 2 * pad,
            QPen(Qt.NoPen), QBrush(Qt.white))
        wr.setZValue(1)
        item = ConstrainedTextItem(self)
        item.setFlag(QGraphicsItem.ItemSendsGeometryChanges, True)
        item.setPlainText(text)
        f = QFont(_map_font(pdf_font))
        sz = 16  # always start at 16
        f.setPointSizeF(sz * FONT_SCALE)
        if "bold" in pdf_font.lower():
            f.setBold(True)
        if any(x in pdf_font.lower() for x in ("italic", "oblique")):
            f.setItalic(True)
        item.setFont(f)
        cr = max(30, (color_int >> 16) & 0xFF)
        cg = max(30, (color_int >> 8) & 0xFF)
        cb = max(30, color_int & 0xFF)
        item.setDefaultTextColor(QColor(cr, cg, cb))
        self._update_text_width(item)
        item.setPos(2 + m, scene_rect.top())
        item.setZValue(2)
        item.setData(0, scene_rect)
        item.setData(1, wr)
        item.setFlag(QGraphicsItem.ItemIsMovable, True)
        item.setFlag(QGraphicsItem.ItemIsSelectable, True)
        item.setTextInteractionFlags(Qt.TextEditorInteraction)
        self.scene.addItem(item)
        wr.stackBefore(item)
        item.setFocus()
        self._current_editor = item
        logger.info("فتح محرر نص قابل للتحريك: '%s'", text[:30])

    def _on_sel_changed(self):
        items = self.scene.selectedItems()
        if not items:
            for item in self.scene.items():
                if isinstance(item, ResizableImageItem) and item._crop_mode:
                    item._crop_mode = False
                    item.update()
        txt_items = [i for i in items if isinstance(i, QGraphicsTextItem)]
        if txt_items:
            self._current_editor = txt_items[0]
            sz = txt_items[0].font().pointSizeF() / FONT_SCALE
            self.font_size_combo.blockSignals(True)
            self.font_size_combo.setCurrentText(f"{sz:.0f}")
            self.font_size_combo.blockSignals(False)
        else:
            self.font_size_combo.blockSignals(True)
            self.font_size_combo.setCurrentText("")
            self.font_size_combo.blockSignals(False)
        self._sync_bold_button()
        self._sync_color_button()

    def _toggle_bold(self):
        item = self._current_editor
        if item is None or not isinstance(item, QGraphicsTextItem):
            return
        self._push_undo(item)
        f = item.font()
        f.setBold(not f.bold())
        item.setFont(f)
        item.update()
        self._bold_btn.setChecked(f.bold())
        logger.info("تبديل الخط العريض: %s", f.bold())

    def _sync_bold_button(self):
        item = self._current_editor
        if item is not None and isinstance(item, QGraphicsTextItem):
            self._bold_btn.setChecked(item.font().bold())
        else:
            self._bold_btn.setChecked(False)

    def _sync_color_button(self):
        item = self._current_editor
        if item is not None and isinstance(item, QGraphicsTextItem):
            c = item.defaultTextColor()
            self._color_btn.setStyleSheet(
                "QPushButton { font-weight: bold; font-size: 14pt; "
                "border-style: solid; border-width: 1px 1px 6px 1px; "
                f"border-color: gray gray {c.name()} gray; }}")
        else:
            self._color_btn.setStyleSheet(
                "QPushButton { font-weight: bold; font-size: 14pt; "
                "border-style: solid; border-width: 1px 1px 6px 1px; "
                "border-color: gray gray red gray; }")

    def _pick_color(self):
        item = self._current_editor
        if item is None or not isinstance(item, QGraphicsTextItem):
            return
        color = QColorDialog.getColor(item.defaultTextColor(), self, "اختر لون النص")
        if not color.isValid():
            return
        self._push_undo(item)
        cursor = item.textCursor()
        if cursor.hasSelection():
            fmt = QTextCharFormat()
            fmt.setForeground(color)
            cursor.mergeCharFormat(fmt)
        else:
            item.setDefaultTextColor(color)
        item.update()
        self._sync_color_button()
        logger.info("تغيير لون النص إلى %s", color.name())

    def _pick_highlight(self):
        item = self._current_editor
        if item is None or not isinstance(item, QGraphicsTextItem):
            return
        color = QColorDialog.getColor(QColor("yellow"), self, "اختر لون الخلفية")
        if not color.isValid():
            return
        self._push_undo(item)
        cursor = item.textCursor()
        if cursor.hasSelection():
            fmt = QTextCharFormat()
            fmt.setBackground(color)
            cursor.mergeCharFormat(fmt)
        else:
            fmt = QTextCharFormat()
            fmt.setBackground(color)
            cursor.select(QTextCursor.Document)
            cursor.mergeCharFormat(fmt)
        item.update()
        logger.info("تغيير خلفية النص إلى %s", color.name())

    def _toggle_numbers(self):
        items = [i for i in self.scene.items() if isinstance(i, QGraphicsTextItem)]
        if not items:
            return
        arabic = '٠١٢٣٤٥٦٧٨٩'
        western = '0123456789'
        first = items[0].toPlainText()
        has_arabic = any(c in arabic for c in first)
        has_western = any(c in western for c in first)
        if has_arabic and not has_western:
            to_arabic = False
        elif has_western and not has_arabic:
            to_arabic = True
        else:
            to_arabic = not self._num_btn.isChecked()
        table = str.maketrans(western, arabic) if to_arabic else str.maketrans(arabic, western)
        for item in items:
            self._push_undo(item)
            item.setPlainText(item.toPlainText().translate(table))
        self._num_btn.setChecked(to_arabic)
        logger.info("تحويل الأرقام: %s", "عربية" if to_arabic else "إنجليزية")

    def _add_image_from_file(self):
        path, _ = QFileDialog.getOpenFileName(self, "إضافة صورة", "",
                                              "Images (*.png *.jpg *.jpeg *.bmp *.gif *.tiff *.webp)")
        if path:
            self._add_image(path)

    def _copy_selected(self):
        items = self.scene.selectedItems()
        for item in items:
            if isinstance(item, ResizableImageItem):
                QApplication.clipboard().setPixmap(item._pix)
                self._copy_pos = item.pos()
                self._copy_rect = QRectF(item._rect)
                logger.info("تم نسخ الصورة")
                return

    def _add_image(self, path):
        pix = QPixmap(path)
        if pix.isNull():
            logger.warning("فشل تحميل الصورة: %s", path)
            return
        page_y = self._current_page_y()
        item = ResizableImageItem(pix)
        # initial size = 100x100 mm
        pw, ph = pix.width(), pix.height()
        target = 100.0
        s = target / max(pw, ph)
        item._rect = QRectF(0, 0, pw * s, ph * s)
        item.setPos(2 + self._margin_mm, page_y + 2)
        item.setZValue(5)
        self.scene.addItem(item)
        logger.info("تم إضافة صورة: %s", os.path.basename(path))

    def _paste_image(self):
        pix = QApplication.clipboard().pixmap()
        if pix.isNull():
            return
        page_y = self._current_page_y()
        item = ResizableImageItem(pix)
        if self._copy_pos is not None and self._copy_rect is not None:
            src_page_y = (self._copy_pos.y() // A4_H) * A4_H
            rel_y = self._copy_pos.y() - src_page_y
            item._rect = QRectF(self._copy_rect)
            item.setPos(self._copy_pos.x() + 10, page_y + rel_y + 10)
        else:
            pw, ph = pix.width(), pix.height()
            target = 100.0
            s = target / max(pw, ph)
            item._rect = QRectF(0, 0, pw * s, ph * s)
            item.setPos(2 + self._margin_mm, page_y + 2)
        item.setZValue(5)
        self.scene.addItem(item)
        self.scene.clearSelection()
        item.setSelected(True)
        logger.info("تم لصق الصورة")

    def _delete_selected(self):
        items = self.scene.selectedItems()
        for item in items:
            if isinstance(item, ResizableImageItem):
                self.scene.removeItem(item)
                logger.info("تم حذف الصورة")
            elif isinstance(item, (ConstrainedTextItem, QGraphicsTextItem)):
                if item is self._current_editor:
                    self._current_editor = None
                self.scene.removeItem(item)
                logger.info("تم حذف النص")

    def _push_undo(self, item=None):
        if item is None:
            item = self._current_editor
        if item is None or not isinstance(item, QGraphicsTextItem):
            return
        state = {
            'item': item,
            'text': item.toPlainText(),
            'font': QFont(item.font()),
            'color': QColor(item.defaultTextColor()),
        }
        self._undo_stack.append(state)
        if len(self._undo_stack) > 50:
            self._undo_stack.pop(0)

    def _undo_last(self):
        if not self._undo_stack:
            return
        state = self._undo_stack.pop()
        item = state['item']
        if item.scene() is None:
            return
        item.setPlainText(state['text'])
        item.setFont(state['font'])
        item.setDefaultTextColor(state['color'])
        self._sync_bold_button()
        logger.info("تراجع عن آخر تعديل")

    def _set_alignment(self, align):
        item = self._current_editor
        if item is None or not isinstance(item, QGraphicsTextItem):
            return
        self._push_undo(item)
        doc = item.document()
        cursor = QTextCursor(doc)
        cursor.select(QTextCursor.Document)
        block_fmt = QTextBlockFormat()
        block_fmt.setAlignment(align)
        cursor.setBlockFormat(block_fmt)
        logger.info("محاذاة النص: %s", align)

    def _font_size_combo_changed(self):
        txt = self.font_size_combo.currentText().strip()
        if not txt:
            return
        try:
            new_sz = float(txt)
        except ValueError:
            return
        item = self._current_editor
        if item is None or not isinstance(item, QGraphicsTextItem):
            return
        self._push_undo(item)
        f = item.font()
        f.setPointSizeF(max(2, new_sz) * FONT_SCALE)
        item.setFont(f)
        self._update_text_width(item)
        item.update()
        logger.info("تغيير حجم الخط إلى %.1f", new_sz)

    def _font_size_up(self):
        self._adjust_font_size(2)

    def _font_size_down(self):
        self._adjust_font_size(-2)

    def _adjust_font_size(self, delta):
        item = self._current_editor
        if item is None or not isinstance(item, QGraphicsTextItem):
            return
        self._push_undo(item)
        f = item.font()
        current_display = f.pointSizeF() / FONT_SCALE
        sz = max(2, current_display + delta)
        f.setPointSizeF(sz * FONT_SCALE)
        item.setFont(f)
        self._update_text_width(item)
        item.update()
        self.font_size_combo.blockSignals(True)
        self.font_size_combo.setCurrentText(f"{sz:.0f}")
        self.font_size_combo.blockSignals(False)
        logger.info("حجم الخط: %.1f", sz)

    def _update_page_label(self):
        scene_top = self.view.mapToScene(0, 0).y()
        page = int(scene_top / A4_H) + 1
        page = max(1, min(page, self._num_pages))
        self._page_combo.blockSignals(True)
        self._page_combo.setEditText(str(page))
        self._page_combo.blockSignals(False)

    def _page_combo_changed(self):
        txt = self._page_combo.currentText().strip()
        try:
            p = int(txt)
        except ValueError:
            return
        p = max(1, min(p, self._num_pages))
        self.view.centerOn(QPointF(0, (p - 1) * A4_H + A4_H / 2))
        logger.info("الانتقال إلى صفحة %d", p)

    def _zoom_in(self):
        prev = self._current_editor
        center = self.view.mapToScene(self.view.viewport().rect().center())
        self.view.scale(1.25, 1.25)
        self.view.centerOn(center)
        if prev is not None:
            prev.setSelected(True)

    def _zoom_out(self):
        prev = self._current_editor
        center = self.view.mapToScene(self.view.viewport().rect().center())
        self.view.scale(1 / 1.25, 1 / 1.25)
        self.view.centerOn(center)
        if prev is not None:
            prev.setSelected(True)

    def _update_text_width(self, item):
        s = item.scale()
        if s > 1e-9:
            item.document().setDocumentMargin(0)
            item.document().setTextWidth((A4_W - 2 * self._margin_mm) / s)

    def _current_page_y(self):
        scene_top = self.view.mapToScene(0, 0).y()
        return (scene_top // A4_H) * A4_H

    def _margin_spin_changed(self, val):
        for item in self.scene.items():
            if isinstance(item, QGraphicsTextItem):
                self._push_undo(item)
        self._margin_mm = max(0, val)
        for item in self.scene.items():
            if isinstance(item, ConstrainedTextItem):
                item.setPos(2 + self._margin_mm, item.y())
                self._update_text_width(item)
                wr = item.data(1)
                if wr is not None:
                    r = wr.rect()
                    wr.setRect(2 + self._margin_mm, r.y(),
                               A4_W - 2 * self._margin_mm, r.height())
        logger.info("تم تطبيق هامش %.1f مم", self._margin_mm)

    def _print_pdf(self):
        if not self._num_pages:
            return
        if not getattr(self, 'subscription_check', lambda: True)():
            return
        from ui.a4_editor import PrintSetupDialog
        from core.printer import print_scene, set_printer_name, set_last_paper_type
        dlg = PrintSetupDialog(self, page_count=self._num_pages, default_paper_type="ورق عادي")
        if dlg.exec() != QDialog.Accepted:
            return
        printer_name = dlg.selected_printer()
        if printer_name:
            set_printer_name(printer_name)
        pt = dlg.paper_type()
        if pt:
            set_last_paper_type(pt)
        logger.info("طباعة %d صفحة", self._num_pages)
        print_scene(self, self.scene, copies=dlg.copies(),
                    page_count=self._num_pages, duplex=dlg.duplex(),
                    page_range=dlg.page_range(), paper_type=pt)

    def _save_as_pdf(self):
        if not self._num_pages:
            return
        if not getattr(self, 'subscription_check', lambda: True)():
            return
        path, _ = QFileDialog.getSaveFileName(self, "حفظ بصيغة PDF", "",
                                              "PDF (*.pdf)")
        if not path:
            return
        if not path.lower().endswith(".pdf"):
            path += ".pdf"
        from PySide6.QtGui import QPdfWriter, QPageSize
        writer = QPdfWriter(path)
        writer.setPageSize(QPageSize(QPageSize.A4))
        writer.setResolution(300)
        painter = QPainter()
        if not painter.begin(writer):
            logger.error("فشل بدء الرسم على ملف PDF عبر QPdfWriter")
            return
        page_rect = painter.viewport()
        scale = min(page_rect.width() / A4_W, page_rect.height() / A4_H)
        pw = A4_W * scale
        ph = A4_H * scale
        ox = (page_rect.width() - pw) / 2
        oy = (page_rect.height() - ph) / 2
        for p in range(1, self._num_pages + 1):
            if p > 1:
                writer.newPage()
            source = QRectF(0, (p - 1) * A4_H, A4_W, A4_H)
            target = QRectF(ox, oy, pw, ph)
            painter.setRenderHint(QPainter.Antialiasing)
            self.scene.render(painter, target, source)
        painter.end()
        if os.path.exists(path) and os.path.getsize(path) > 0:
            logger.info("تم حفظ PDF بنجاح: %s (%d بايت)", path, os.path.getsize(path))
        else:
            logger.error("فشل حفظ PDF: الملف فارغ أو غير موجود: %s", path)

    def _save_as_docx(self):
        if not self._num_pages:
            return
        if not getattr(self, 'subscription_check', lambda: True)():
            return
        path, _ = QFileDialog.getSaveFileName(self, "حفظ بصيغة Word", "",
                                              "Word (*.docx)")
        if not path:
            return
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.oxml.ns import qn
        doc = Document()
        for p in range(self._num_pages):
            if p > 0:
                doc.add_page_break()
            for item in self.scene.items():
                if not isinstance(item, QGraphicsTextItem):
                    continue
                y = item.sceneBoundingRect().top()
                page_y0 = p * A4_H
                page_y1 = (p + 1) * A4_H
                if page_y0 <= y < page_y1:
                    text = item.toPlainText().strip()
                    if not text:
                        continue
                    para = doc.add_paragraph()
                    run = para.add_run(text)
                    f = item.font()
                    run.font.size = Pt(f.pointSizeF() / FONT_SCALE)
                    run.font.bold = f.bold()
                    run.font.italic = f.italic()
                    c = item.defaultTextColor()
                    run.font.color.rgb = RGBColor(c.red(), c.green(), c.blue())
                    rtl = doc.styles['Normal'].element.rPr
                    if rtl is None:
                        from docx.oxml import OxmlElement
                        rPr = OxmlElement('w:rPr')
                        doc.styles['Normal'].element.append(rPr)
                    doc.styles['Normal'].element.rPr.set(qn('w:rtl'), '1')
                    pf = para.paragraph_format
                    pf.space_before = Pt(0)
                    pf.space_after = Pt(0)
        doc.save(path)
        logger.info("تم حفظ Word: %s", path)


_FONT_MAP = {
    "times-roman": "Times New Roman", "timesnewroman": "Times New Roman",
    "timesbold": "Times New Roman", "timesitalic": "Times New Roman",
    "helvetica": "Arial", "arialmt": "Arial", "arial-boldmt": "Arial",
    "courier": "Courier New", "couriernew": "Courier New",
    "symbol": "Symbol", "zapfdingbats": " Wingdings",
}


def _map_font(pdf_name):
    key = pdf_name.replace("-", "").replace(" ", "").lower()
    for k, v in _FONT_MAP.items():
        if k in key:
            return v
    if any(c.isalpha() for c in pdf_name):
        return pdf_name
    return "Arial"


